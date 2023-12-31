import io
import itertools
import re
from difflib import SequenceMatcher
import streamlit as st
from streamlit.components.v1 import html
from docx import Document
from typing import List, Tuple
from fuzzywuzzy import fuzz
from metaphone import doublemetaphone
from streamlit import components
from io import BytesIO
import base64
from pathlib import Path
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import regex
import os

# Ensure preceding_names, succeeding_names, new_text, replaced_names, and unmatched_names are in session state
if 'extracted_names' not in st.session_state:
    st.session_state.extracted_names = []
if 'formatted_names' not in st.session_state:
    st.session_state.formatted_names = []
if 'preceding_names' not in st.session_state:
    st.session_state.preceding_names = []
if 'succeeding_names' not in st.session_state:
    st.session_state.succeeding_names = []
if 'new_text' not in st.session_state:
    st.session_state.new_text = ""
if 'replaced_names' not in st.session_state:
    st.session_state.replaced_names = []
if 'unmatched_names' not in st.session_state:
    st.session_state.unmatched_names = []

american_to_british_dict = {
  'honored':'honoured',
  'honor':'honour',
  'realizing':'realising',
  'realize':'realise',
  'color':'colour',
  'colored':'coloured',
  'recognize':'recognise',
  'recognizes':'recognises',
  'recognized':'recognised',
  'humor':'humour',
  'humored':'humoured',
  'organize':'organise',
  'organized':'organised',
  'eon':'aeon',
  'esthetic':'aesthetic',
  'anemia':'anaemia',
  'anesthesia':'anaesthesia',
  'pediatrician':'paediatrician',
  'appall':'appal',
  'carburetor':'carburettor',
  'counselor':'counsellor',
  'disheveled':'dishevelled',
  'distill':'distil',
  'enroll':'enrol',
  'fufill':'fulfil',
  'installment':'instalment',
  'instill':'instil',
  'skillful':'skilful',
  'woollen':'woollen',
  'defense':'defence',
  'license':'licence',
  'offense':'offence',
  'pretense':'pretence',
  'annex':'annexe',
  'glycerin':'glycerine',
  'gram':'gramme',
  'program':'programme',
  'ton':'tonne',
  'diarrhea':'diarrhoea',
  'maneuver':'manoeuvre',
  'arbor':'arbour',
  'armor':'armour',
  'behavior':'behaviour',
  'candor':'candour',
  'clamor':'clamour',
  'demeanor':'demeanour',
  'endeavor':'endeavour',
  'favor':'favour',
  'flavor':'flavour',
  'habor':'harbour',
  'labor':'labour',
  'neighbor':'neighbour',
  'odor':'odour',
  'parlor':'parlour',
  'rancor':'rancour',
  'rigor':'rigour',
  'rumor':'rumour',
  'savior':'saviour',
  'savor':'savour',
  'splendor':'splendour',
  'tumor':'tumour',
  'valor':'valour',
  'vigor':'vigour',
  'caliber':'calibre',
  'center':'centre',
  'fiber':'fibre',
  'liter':'litre',
  'luster':'lustre',
  'meager':'meagre',
  'meter':'metre',
  'saber':'sabre',
  'scepter':'sceptre',
  'sepulcher':'sepulchre',
  'somber':'sombre',
  'theater':'theatre',
  'airplane':'aeroplane',
  'artifact':'artefact',
  'checkerboard':'chequerboard',
  'checkered':'chequered',
  'cozy':'cosy',
  'donut':'doughnut',
  'draft':'draught',
  'jewelry':'jewellery',
  'curb':'kerb',
  'plow':'plough',
  'judgment':'judgement',
  'civilize':'civilise',
  'analyze':'analyse',
  'paralyze':'paralyse',
  'humorous':'humourous',
  'civilization':'civilisation',
  'acknowledgment':'acknowledgement',
  'favorite':'favourite',
  'benefited':'benefitted',
  'practise':'practice',
  'characterization':'characterisation',
  'learned':'learnt',
  'decarbonization ':'decarbonisation',
  'dependent':'dependant',
  'judgmental':'judgemental',
  'apologize':'apologise',
  'fulfill':'fulfil',
}

# Name Extractor for graduation ceremony in-person lists functions
def extract_names(doc):
    middle_column_texts = []

    # List of phrases to exclude
    excluded_phrases = ["vacant seat", "carer's seat", "carers seat", "child", "seat for pa companion", 
                        "pa companion", "pa companion seat", "companion seat", "pa companion's seat", "companion's seat"]

    # Compile all excluded phrases into a single regular expression
    excluded_phrases_regex = re.compile("|".join(excluded_phrases), re.IGNORECASE)

    # Regular expression for name extraction
    name_regex = re.compile(r'\b[A-Z][a-zA-Z]*\b')

    # Iterate over all tables and rows
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) > 1:
                # Grab the first and middle cell
                first_cell = cells[0]
                middle_cell = cells[len(cells) // 2]

                # Get all the paragraphs in the middle cell
                paragraphs = middle_cell.paragraphs

                for paragraph in paragraphs:
                    text = paragraph.text.strip()

                    # Check if the text is strikethrough
                    is_strikethrough = any(run.font.strike for run in paragraph.runs)

                    words = text.split()

                    # Adjust word count if "(Marked As Not Present)" is in text
                    word_count = len(words) - 4 if is_strikethrough else len(words)

                    # Check if the text is not in excluded phrases and:
                    # 1. Starts with an uppercase word and the line of text is not longer than six words, or
                    # 2. Ends with an uppercase word and the line of text is not longer than six words, or
                    # 3. Matches a name pattern and the first cell is empty, or
                    # 4. There is only a single line or paragraph in the middle column cell
                    if not excluded_phrases_regex.search(text) and \
                       ((re.search(r'^\b[A-Z]+\b', text) and word_count <= 7) or 
                        (re.search(r'\b[A-Z]+\b$', text) and word_count <= 7) or 
                        (name_regex.search(text) and not first_cell.text.strip()) or
                        (len(paragraphs) == 1)):
                        # This line contains the name. 
                        # Add note if name is strikethrough
                        if is_strikethrough:
                            text += ' (Marked As Not Present)'

                        # Remove single-letter words (accounting for periods and spaces)
                        words = re.split("[ .]", text)  # split on spaces and periods
                        text = ' '.join(word for word in words if len(word) > 1)

                        middle_column_texts.append(text)

    # Decapitalize names before returning
    middle_column_texts = [decapitalize(name) for name in middle_column_texts]

    return middle_column_texts
  
def format_names(names_list):
    # Add more colors if needed
    colors = ['red', 'green', 'blue', 'yellow']
    formatted_names = []
    for i, name in enumerate(names_list):
        color = colors[i % len(colors)]
        formatted_name = (name, color)
        formatted_names.append(formatted_name)
    return formatted_names

def decapitalize(text):
    roman_numerals = ['I', 'II', 'III', 'IV', 'V', 'VI']
    words = text.split()
    for i, word in enumerate(words):
        if word not in roman_numerals:

            # Split hyphenated words and capitalize each part
            hyphen_parts = word.split('-')
            hyphen_parts = [part.lower().title() for part in hyphen_parts]
            word = '-'.join(hyphen_parts)

            # Split words with apostrophes and capitalize each part
            apostrophe_parts = word.split("'")
            apostrophe_parts = [part.lower().title() for part in apostrophe_parts]
            word = "'".join(apostrophe_parts)

            # If a name starts with 'Mc', capitalize the next letter
            if word.startswith('Mc') and len(word) >= 3:
                word = word[:2] + word[2].upper() + word[3:]

            words[i] = word

    return ' '.join(words)

#Correct all names in graduation subtitles (find and replace) functions

    # Create a regex pattern to match the name with optional whitespace between words
    pattern = r'\b' + r'\s*'.join(re.escape(word) for word in name_words) + r'\b'

    # Find all occurrences of the pattern in the text (case-insensitive)
    similar_names = re.findall(pattern, text, flags=re.IGNORECASE)

    # Remove duplicates and return the result
    return list(set(similar_names))

def similarity(a, b):
    sequence_similarity = SequenceMatcher(None, a, b).ratio()
    fuzz_similarity = fuzz.ratio(a, b) / 100.0
    metaphone_similarity = doublemetaphone(a) == doublemetaphone(b)

    # Calculate overall similarity
    overall_similarity = sequence_weight * sequence_similarity + fuzz_weight * fuzz_similarity + metaphone_weight * metaphone_similarity

    # Apply penalty if a name in the subtitles (a) has less words than a name in the list (b)
    if len(a.split()) < len(b.split()):
        penalty_factor = 0.05  # Adjust this value to increase or decrease the penalty
        overall_similarity *= penalty_factor

    return overall_similarity

def replace_similar_names(text: str, names_list: List[str]) -> Tuple[List[Tuple[str, str, float]], str]:
    replaced_names = []
    unmatched_names = names_list[:]  # Make a copy of names_list

    def replace_name(match):
        full_name = match.group(0)
        # Check if the name is already replaced
        for original, replaced, _ in replaced_names:
            if full_name == replaced:
                return full_name
    
        max_similarity = 0
        most_similar_name = None
        for name in names_list:
            # Remove the "(Marked As Not Present)" marker for comparison
            clean_name = name.replace(' (Marked As Not Present)', '')
            sim = similarity(full_name, clean_name)
            if sim > max_similarity and (not match_word_count or len(full_name.split()) == len(clean_name.split())):  # Compare with clean_name
                max_similarity = sim
                most_similar_name = clean_name  # Use clean_name to replace
    
        if max_similarity >= similarity_threshold:
            replaced_names.append((full_name, most_similar_name, max_similarity))
            # Remove the name from unmatched_names if it was matched
            if most_similar_name in unmatched_names:
                unmatched_names.remove(most_similar_name)
            return most_similar_name
        else:
            return full_name

    # Updated regex pattern
    pattern = r'\b([A-Z][a-z]+(?:(?: |-)[A-Z][a-z]+)*)\b'

    processed_lines = []
    lines = text.split('\n')
    for line in lines:
        # Skip timecode lines
        if re.match(r'\d\d:\d\d:\d\d\.\d\d\d\s*-->', line):
            processed_lines.append(line)
            continue

        line = re.sub(pattern, replace_name, line)
        processed_lines.append(line)

    new_text = '\n'.join(processed_lines)

    if replaced_names:
        # Remove leading whitespaces from all lines as a final step
        new_text = '\n'.join(line.lstrip() for line in new_text.split('\n'))
        return replaced_names, new_text, unmatched_names  # Return unmatched_names as well
    else:
        return [], '', unmatched_names  # Return unmatched_names as well

def reformat_subtitles(text: str) -> str:
    if text.startswith('WEBVTT'):
        text = text.replace('WEBVTT', 'WEBVTT\n', 1)

    blocks = re.split(r'((?:\d\d:)?\d\d:\d\d\.\d\d\d --> (?:\d\d:)?\d\d:\d\d\.\d\d\d)', text)
    formatted_blocks = []

    for block in blocks:
        if re.match(r'(?:\d\d:)?\d\d:\d\d\.\d\d\d --> (?:\d\d:)?\d\d:\d\d\.\d\d\d', block):
            formatted_blocks.append(block.strip() + '\n')
            continue

        lines = block.split('\n')
        lines = [line for line in lines if line.strip()]

        formatted_lines = []
        for line in lines:
            words = line.split()
            if words:
                formatted_line = ' '.join(words)
            formatted_line = re.sub(r'\[no audio\]', '', formatted_line, flags=re.IGNORECASE)
            formatted_line = re.sub(r'\((applause|ALL APPLAUD|APPLAUSE CONTINUES)\)|\[(applause|ALL APPLAUD|APPLAUSE CONTINUES)\]|Applause|\*\* Applause \*\*|👏👏👏|👏|ALL APPLAUD|APPLAUSE CONTINUES', '[Audience Applauds]', formatted_line, flags=re.IGNORECASE)
            formatted_line = re.sub(r'\((Music|MUSIC|MUSIC PLAYING|ORGAN MUSIC|ORGAN MUSIC CONTINUES|ORCHESTRAL MUSIC| Music )\)|\[(Music|MUSIC|MUSIC PLAYING|ORCHESTRAL MUSIC| Music )\]|♪♪|MUSIC PLAYING|ORGAN MUSIC|ORGAN MUSIC CONTINUES|ORCHESTRAL MUSIC', '[Music Playing]', formatted_line, flags=re.IGNORECASE)
            formatted_line = re.sub(r'\((laughter|ALL LAUGH)\)|\[(laughter|ALL LAUGH)\]|laughter|ALL LAUGH', '[Audience Laughing]', formatted_line, flags=re.IGNORECASE)
            formatted_line = re.sub(r'\((cheering|audience cheering|CHEERING AND APPLAUSE|INDISTINCT CHATTER)\)|\[(cheering|audience cheering|CHEERING AND APPLAUSE|INDISTINCT CHATTER)\]', '[Audience Cheers]', formatted_line, flags=re.IGNORECASE)
            formatted_line = re.sub(r'\((shouting|audience shouting)\)|\[(shouting|audience shouting)\]|audience shouting', '[Audience Shouts]', formatted_line, flags=re.IGNORECASE)
            formatted_line = re.sub(r'\((pause)\)|\[(pause)\]', '[Pause]', formatted_line, flags=re.IGNORECASE)
            formatted_line = re.sub(r'\((exhale)\)|\[(exhale)\]', '[They Exhale]', formatted_line, flags=re.IGNORECASE)

            # Convert dict keys/values to lowercase
            local_american_to_british_dict = {k.lower(): v.lower() for k, v in american_to_british_dict.items()}

            # American to British replacement
            for american, british in local_american_to_british_dict.items():
                # Replace whole words only
                formatted_line = re.sub(rf'\b{american}\b', british, formatted_line, flags=re.IGNORECASE)

            formatted_lines.append(formatted_line)

        formatted_block = '\n'.join(formatted_lines) + '\n\n'
        formatted_blocks.append(formatted_block)

    return ''.join(formatted_blocks)  # Return as a string

def reformat_transcript(text: str):
    # Remove 'WEBVTT'
    text = text.replace('WEBVTT', '')

    # Remove timestamps, and empty lines
    text = re.sub(r'(\d\d:\d\d:\d\d\.\d\d\d\s-->\s\d\d:\d\d:\d\d\.\d\d\d|\d\d:\d\d\.\d\d\d\s-->\s\d\d:\d\d\.\d\d\d)\n?', '', text)
    
    # Replace paragraph breaks and multiple spaces with single spaces
    formatted_text = re.sub(r'\s+', ' ', text)
    
    # Remove anything in square or round brackets
    formatted_text = re.sub(r'\[.*?\]', '', formatted_text)  # Square brackets
    formatted_text = re.sub(r'\(.*?\)', '', formatted_text)  # Round brackets

    # Write to Word file
    doc = Document()

    # Add initial text
    p = doc.add_paragraph()
    p.add_run('[‘Trumpet Fanfare’ music playing] (A procession of University senior academics and staff in ceremonial robes enter the auditorium, walk down the aisles betwixt the audience of seated graduands and guests, ascend the stage via staircases on the left and right respectively, and take their seats. At the end of the procession are two academics/staff with ceremonial torches who on stage bow to each other, the rows of academics/staff, and then place the torches on a small, raised table with a cloth at the very front of the stage.)').italic = True
    doc.add_paragraph()  # Add paragraph break

    # Add the formatted transcript
    doc.add_paragraph(formatted_text)

    # Add final text
    p = doc.add_paragraph()
    p.add_run('[Music playing] (Senior academics and staff on stage tip their hats as two academics/staff walk across the stage to pick up the ceremonial torches from the small, raised table. They bow to one another before bowing to the rest of the academics/staff. Both lead lines single file of all the professors in separate directions down the staircases on the left and right. The academics and staff walk down the aisles betwixt the audience of seated graduates and guests and exit at the back of the auditorium.)').italic = True

    # Write to a BytesIO buffer instead of a file
    buffer = BytesIO()
    doc.save(buffer)

    # Return the buffer along with the formatted text
    return formatted_text, buffer

def get_binary_file_downloader_html(bin_file, file_label='File'):
    with open(bin_file, 'rb') as f:
        data = f.read()
    bin_str = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{Path(bin_file).name}">{file_label}</a>'
    return href
        
#Name Corrector UI
st.markdown("""
<h1 style='text-align: center; font-size: 60px; color: white;'>Graduation Transcription Workflow Web Tool</h1>
""", unsafe_allow_html=True)
st.markdown("[Read the Workflow Document here for more information on how to use this app.](https://sussex.box.com/s/0mkoals6oqy9tkiasshos917zrlorrha)")

# Add a slider in the sidebar
st.sidebar.header('Set Overall Similarity Threshold for Combined Methods')
similarity_threshold = st.sidebar.slider(
    'Set your similarity threshold. Lower values make name matching more lenient, higher values make it stricter. Start at 0.30, identify when nonsense matches begin in Names Replaced, increase the threshold and run again.',
    min_value=0.0,
    max_value=1.0,
    value=0.30,
    step=0.01,
)

# Slider weights
st.sidebar.header('Adjust Weights for Comparison Methods')
st.sidebar.text('Set the relative weights of each method towards the name similarity matching - experimental.')
sequence_weight = st.sidebar.slider ('SequenceMatcher Weight', 0.0, 1.0, 0.33, 0.01)
fuzz_weight = st.sidebar.slider ('Fuzz Ratio Weight', 0.0, 1.0, 0.33, 0.01)
metaphone_weight = st.sidebar.slider ('Double Metaphone Weight', 0.0, 1.0, 0.34, 0.01)

# Ensure the sum of weights equal to 1
if sequence_weight + fuzz_weight + metaphone_weight != 1.0:
    st.sidebar.error('Please adjust the weights so their sum equals to 1.0')

# Add the banner image at the top of the app
st.image("banner2.jpg")

# Match Word Count UI
st.sidebar.header('Match Word Count')
st.sidebar.text('Turning on ensures less mismatching, but more necessary if only relying on SequenceMatcher.')
match_word_count = st.sidebar.checkbox('Should the number of words match?', value=False)

with st.expander("1 - Follow URL Below To Generate Subtitle VTT File From Vimeo URL"):

    # Define the URL
    url = "https://colab.research.google.com/drive/1mGeZ_2qnc8KrLyV-U1xR5wtvf6MV8zFa"
    
    # Define the image URL
    image_url = "https://github.com/Alphadoriest/UniversityofSussex/blob/Main-6/Collab.png?raw=true"  # replace with your image URL
    
    # Create the HTML string
    html = f'<a href="{url}" target="_blank"><img src="{image_url}" alt="Colab link" style="width:853px; height:423px;"/></a>'
    
    # Display the HTML in the app
    st.markdown(html, unsafe_allow_html=True)

with st.expander("2 - Name Extractor for Graduation Ceremony In-Person Lists"):

    uploaded_file = st.file_uploader("Choose a Ceremony In-Person List Word document", type="docx")
    
    # Initialize data as an empty list
    data = []
    
    if uploaded_file is not None:
        document = Document(io.BytesIO(uploaded_file.read()))
        data = extract_names(document)  # Keep data as a list
    
    # Use data (the names_list) directly, as there is no 'Name' key to extract from
    names_list = data
    
    # Use names_list as the default value for the names_list text_area
    names_list = [name for name in names_list if name is not None]
    names_list = st.text_area("Alternatively, enter names, separated by commas:", ', '.join(names_list), key='names_list')
    
    if names_list:  # Check if names_list is not empty
        names_list = names_list.split(',')
        names_list = [name.strip() for name in names_list]

    # Remove single-letter words from each name and apply decapitalize
        names_list = [' '.join(word for word in name.split() if len(word) > 1) for name in names_list if name not in ["VACANT SEAT", "Vacant Seat", "Carer's seat", "CARER'S SEAT", "Child", "CHILD","Seat for PA Companion", "PA Companion", "PA Companion seat", "Companion Seat",]]
        names_list = [decapitalize(name) for name in names_list]
        
        # Check if names_list contains meaningful entries
        if any(name for name in names_list):
            # Assuming format_names now returns a list of tuples like [(name, color), ...]
            formatted_names = format_names(names_list)
        
        # Create the names list as a Markdown string
        names_md = ', '.join(
            [
                f'<span style="color:{color};"><strong><u>{name}</u></strong></span>'
                if (len(name.split()) > 7 or re.search(r'[()\[\]{}\",;&*=+/?:.]', name))
                else f'<span style="color:{color};">{name}</span>'
                for name, color in formatted_names
            ]
        )
        st.text("If a name contains more than 7 words or contains brackets the name is underlined and made bold so it's easy to spot potential errors in extraction.")    
        # Display the names list using st.markdown
        st.markdown(names_md, unsafe_allow_html=True)
            
# Create a collapsible section or container for the Graduation Subtitles Name Corrector
with st.expander("3 - Graduation Subtitles Name Corrector"):

            # Initialize subtitles_text as an empty string
            subtitles_text = ''

            uploaded_subtitles_file = st.file_uploader("Choose a VTT or TXT Subtitles File ", type=["vtt", "txt"])
            if uploaded_subtitles_file is not None:
                subtitles_text = uploaded_subtitles_file.read().decode()

            # Use the subtitles_text as the default value for the subtitles text_area
            text = st.text_area("Alternatively, Enter Text From a Subtitles:", subtitles_text, key='subtitles_text')

            # Add a separate button for the name replacement process
            if st.button("Press to Replace Names"):  
                if names_list and text:  # Check if both text boxes are populated
                    replaced_names, new_text, unmatched_names = replace_similar_names(text, names_list)  # Unpack unmatched_names

                    # Store the resultant text and replaced_names and unmatched_names in session state
                    st.session_state.new_text = reformat_subtitles(new_text)  # Use reformat_subtitles here
                    st.session_state.replaced_names = replaced_names
                    st.session_state.unmatched_names = unmatched_names
                    # Get the indices of unmatched names in names_list
                    unmatched_indices = [names_list.index(name) for name in st.session_state.unmatched_names if name in names_list]
                    # Get the names that precede the unmatched names and store in the session state
                    st.session_state.preceding_names = [names_list[i-1] if i > 0 else None for i in unmatched_indices]
                    # Get the names that succeed the unmatched names and store in the session state
                    st.session_state.succeeding_names = [names_list[i+1] if i < len(names_list) - 1 else None for i in unmatched_indices]
            
            # Display replaced, unmatched, preceding, and succeeding names from session state
            st.subheader("Names replaced:")
            for original, replaced, similarity in sorted(st.session_state.replaced_names, key=lambda x: -x[2]):  # Sort by similarity
                original_words = original.split()
                replaced_words = replaced.split()
                if len(original_words) != len(replaced_words):
                    st.markdown(f"**{original} -> {replaced} (Similarity: {similarity:.2f})**")
                else:
                    st.write(f"{original} -> {replaced} (Similarity: {similarity:.2f})")
            
            st.subheader("Names not matched:")
            st.text("These can be addressed in one of two ways. Either copy the comma separated list and run just those names in another instance of the app at a lower threshold or browser search for the names surrounding the unmatched name and paste in the correct name in the updated subtitles text box. The app will reset after each addition, but all progress is saved.")
            unmatched_names_str = ', '.join(st.session_state.unmatched_names)
            st.write(unmatched_names_str)
            
            # Button to copy unmatched names to clipboard
            copy_unmatched_names_button_html = f"""
                <button onclick="copyUnmatchedNames()">Copy unmatched names to clipboard</button>
                <script>
                function copyUnmatchedNames() {{
                    navigator.clipboard.writeText("{unmatched_names_str}");
                }}
                </script>
                """
            components.v1.html(copy_unmatched_names_button_html, height=30)
            
            # Get the indices of unmatched names in names_list
            unmatched_indices = [names_list.index(name) for name in st.session_state.unmatched_names if name in names_list]
            
            # Get the names that precede the unmatched names
            preceding_names = [names_list[i-1] if i > 0 else None for i in unmatched_indices]
            
            # Get the names that succeed the unmatched names
            succeeding_names = [names_list[i+1] if i < len(names_list) - 1 else None for i in unmatched_indices]
            
            st.subheader("Preceding and Succeeding Names for Easy Look Up of Unmatched Name for Addition to Updated Subtitles Box:")
            for preceding, succeeding, unmatched in zip(st.session_state.preceding_names, st.session_state.succeeding_names, st.session_state.unmatched_names):
                st.write(f"{preceding or 'N/A'}, {succeeding or 'N/A'} -> {unmatched}")
            
            # Get the text from the text area
            new_text = st.text_area("Updated Subtitles Text to Copy Into VTT/TXT File:", st.session_state.get('new_text', ''), key='updated_subtitles_text')
            
            # Save changes button
            if st.button('Save Changes'):
                # Update session state with any changes made in the text area
                st.session_state.new_text = reformat_subtitles(new_text)  # Use reformat_subtitles here
            
            st.markdown("To copy the replaced text to the clipboard, manually select the text above and use your browser's copy function (right-click and select 'Copy' or use the keyboard shortcut Ctrl/Cmd+C).")        

with st.expander("4 - Reformat Your VTT Into a Word Transcript"):
    # Initialize transcript_text as an empty string
    transcript_text = ''
    original_file_name = ''  # Initialize to ensure variable exists even if no file is uploaded
    
    uploaded_transcript_file = st.file_uploader("Choose a Transcript VTT or TXT File ", type=["vtt", "txt"])
    if uploaded_transcript_file is not None:
        original_file_name, _ = os.path.splitext(uploaded_transcript_file.name)
        transcript_text = uploaded_transcript_file.read().decode()
    
    # Use the transcript_text as the default value for the transcript text_area
    transcript_text = st.text_area("Alternatively, Enter VTT/TXT Text:", transcript_text, key='transcript_text')
    
    # Reformat the transcript when a button is pressed
    if st.button("Reformat VTT/TXT Into Transcript", key="reformat_button"):
        if transcript_text:  # Check if transcript_text is not empty
            reformatted_transcript, buffer = reformat_transcript(transcript_text)  # Unpack buffer
            transcript_text = reformatted_transcript  # Overwrite transcript_text with the reformatted transcript    
        
        # Display the reformatted transcript
        st.text_area("Reformatted Transcript:", transcript_text, key='reformatted_transcript')
        
        # Provide download link for the Word file
        buffer.seek(0)  # Reset buffer position
        download_file_name = f"{original_file_name} Transcript.docx"
        st.download_button('Download Word file', buffer, download_file_name, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')

with st.expander("5 - Use the Closed Captions Pool to add Captions to the Reformatted Transcript"):
    st.markdown("[Closed Captions Pool](https://sussex.box.com/s/3asyl10xocjs0qmdkzpr39mvz9kukdsb)")
